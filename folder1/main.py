import os
import streamlit as st
st.set_page_config(page_title="Research Tool", layout="wide")

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from dotenv import load_dotenv
import time
import requests
import re
from datetime import datetime, timedelta, date
import yake
from collections import Counter
import PyPDF2
import io
import docx
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder
import numpy as np
import cohere
from advanced_features import (
    MetadataFilter, filtered_similarity_search,
    FeedbackStorage,
    ConversationMemory, AdaptiveContextManager,
    Crawl4AIMetadataExtractor,
    HybridMetadataExtractor,
    WebMetadataQualityAssessor,
    MetadataEnricher,
    ContentQualityFilter,
    CredibilityScorer,
    DomainSpecificFilter,
    AdvancedMetadataFilter
)

from conversation_continuity import ConversationContinuitySystem

load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")
news_api_key = os.getenv("NEWS_API_KEY", "c8ad3a77b1aa46a68541fe8a9b56ac7f")
cohere_api_key = os.getenv("COHERE_API_KEY", "")  

if not openai_api_key:
    st.error("❌ OpenAI API key not found! Please add your OPENAI_API_KEY to the .env file")
    st.stop()

def set_background_color():
    st.markdown(
        """
        <style>
        .stApp {
            background-color: #76ABAE;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
set_background_color()

if 'metadata_filter' not in st.session_state:
    st.session_state.metadata_filter = MetadataFilter()
if 'feedback_storage' not in st.session_state:
    st.session_state.feedback_storage = FeedbackStorage("feedback.db")
if 'conversation_memory' not in st.session_state:
    st.session_state.conversation_memory = ConversationMemory(max_history=10)
if 'adaptive_context' not in st.session_state:
    st.session_state.adaptive_context = AdaptiveContextManager(st.session_state.conversation_memory)
if 'session_id' not in st.session_state:
    st.session_state.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")

if 'continuity_system' not in st.session_state:
    st.session_state.continuity_system = ConversationContinuitySystem()

st.title("Research Tool")
st.caption("With Metadata Filtering, Feedback Loops & Adaptive Context")


app_tab1, app_tab2, app_tab3, app_tab4, app_tab5 = st.tabs(["💬 Chat", "📊 Analytics", "⚙️ Settings", "🔍 Metadata Extraction", "🔄 News/Conversation Track"])

with app_tab1:
    tab1, tab2, tab3 = st.tabs(["🌐 URLs", "📝 Direct Text", "📁 File Upload"])

st.sidebar.header("📰 News Settings")
enable_news = st.sidebar.checkbox("Enable News Articles", value=True)
num_news_articles = st.sidebar.slider("Number of News Articles", 1, 10, 5)
show_keywords = st.sidebar.checkbox("Show Extracted Keywords", value=False)
debug_yake = st.sidebar.checkbox("Debug YAKE Output", value=False)

st.sidebar.header("🎯 Re-Ranking Settings")
enable_reranking = st.sidebar.checkbox("Enable Advanced Re-Ranking", value=True)
reranking_strategy = st.sidebar.selectbox(
    "Re-Ranking Strategy",
    [
        "Ultimate Hybrid (All Methods + RRF)",
        "Hybrid (BM25 + Semantic + Cross-Encoder)", 
        "Cohere Re-rank (Industry-leading)",
        "RankGPT (LLM-based)",
        "BM25 Only", 
        "Semantic Only", 
        "Cross-Encoder Only"
    ],
    index=0
)
num_candidates = st.sidebar.slider("Number of candidates to retrieve", 4, 20, 10)
num_final_results = st.sidebar.slider("Number of final results after re-ranking", 2, 10, 4)
show_confidence_scores = st.sidebar.checkbox("Show Confidence Scores", value=True)

st.sidebar.header("🔍 Debug Settings")
enable_debug = st.sidebar.checkbox("Enable Debug Mode", value=False)
show_chunks = st.sidebar.checkbox("Show Retrieved Chunks", value=False)
show_processing_steps = st.sidebar.checkbox("Show Processing Steps", value=False)
show_storage_info = st.sidebar.checkbox("Show Storage Information", value=False)
st.sidebar.header("🔎 Metadata Filters")
st.sidebar.caption("Filter documents before retrieval")
if st.sidebar.button("🗑️ Clear All Filters"):
    st.session_state.metadata_filter.clear()
    st.rerun()
use_date_filter = st.sidebar.checkbox("📅 Filter by Date Range")
if use_date_filter:
    col1, col2 = st.sidebar.columns(2)
    with col1:
        start_date = st.date_input("From", value=date.today() - timedelta(days=365), key="filter_start_date")
    with col2:
        end_date = st.date_input("To", value=date.today(), key="filter_end_date")
    
    st.session_state.metadata_filter.add_date_range_filter(
        start_date=start_date.isoformat(),
        end_date=end_date.isoformat()
    )
use_source_filter = st.sidebar.checkbox("📁 Filter by Source Type")
if use_source_filter:
    source_types = st.sidebar.multiselect(
        "Source Types",
        ["url", "pdf", "docx", "txt"],
        default=["url", "pdf", "docx", "txt"],
        key="filter_source_types"
    )
    if source_types:
        st.session_state.metadata_filter.add_source_type_filter(source_types)

use_author_filter = st.sidebar.checkbox("👤 Filter by Author")
if use_author_filter:
    authors_input = st.sidebar.text_input("Authors (comma-separated)", key="filter_authors")
    if authors_input:
        author_list = [a.strip() for a in authors_input.split(",") if a.strip()]
        if author_list:
            st.session_state.metadata_filter.add_author_filter(author_list)
use_category_filter = st.sidebar.checkbox("📂 Filter by Category")
if use_category_filter:
    categories = st.sidebar.multiselect(
        "Categories",
        ["technology", "business", "science", "health", "politics", "education", "other"],
        key="filter_categories"
    )
    if categories:
        st.session_state.metadata_filter.add_category_filter(categories)

use_tag_filter = st.sidebar.checkbox("🏷️ Filter by Tags")
if use_tag_filter:
    tags_input = st.sidebar.text_input("Tags (comma-separated)", key="filter_tags")
    match_all_tags = st.sidebar.checkbox("Match ALL tags (default: match ANY)", key="filter_match_all")
    if tags_input:
        tag_list = [t.strip() for t in tags_input.split(",") if t.strip()]
        if tag_list:
            st.session_state.metadata_filter.add_tag_filter(tag_list, match_all=match_all_tags)
if st.session_state.metadata_filter.filters:
    st.sidebar.success(f"✅ {len(st.session_state.metadata_filter.filters)} filter(s) active")
    with st.sidebar.expander("View Active Filters"):
        for key, value in st.session_state.metadata_filter.filters.items():
            if isinstance(value, datetime):
                st.sidebar.text(f"{key}: {value.date()}")
            else:
                st.sidebar.text(f"{key}: {value}")
st.sidebar.header("💬 Conversation")
if st.session_state.conversation_memory.history:
    st.sidebar.write(st.session_state.conversation_memory.get_summary())
    
    if st.sidebar.button("🗑️ Clear History"):
        st.session_state.conversation_memory.clear()
        st.rerun()
    
    with st.sidebar.expander("📜 View History"):
        for i, turn in enumerate(reversed(st.session_state.conversation_memory.history)):
            st.sidebar.markdown(f"**Turn {len(st.session_state.conversation_memory.history) - i}**")
            st.sidebar.text(f"Q: {turn['query'][:50]}...")
            st.sidebar.text(f"Confidence: {turn['confidence']:.0f}%")
            st.sidebar.divider()
else:
    st.sidebar.info("No conversation history yet")

start_processing = False
input_data = None
data_type = None

with tab1:
    st.header("🌐 URLs")
    st.write("Enter up to 3 URLs to process and ask questions about their content.")
    
    input_urls = [st.text_input(f"URL {i+1}", key=f"url_{i}") for i in range(3)]
    start_processing = st.button("🚀 Process URLs", key="process_urls")
    
    if start_processing:
        valid_urls = [url for url in input_urls if url.strip()]
        if valid_urls:
            input_data = valid_urls
            data_type = "urls"
        else:
            st.error("Please enter at least one valid URL")

with tab2:
    st.header(" Direct Text Input")
    st.write("Paste or type your text content directly and ask questions about it.")
    
    text_input = st.text_area(
        "Enter your text content:",
        height=300,
        placeholder="Paste your text here...",
        key="text_input"
    )
    
    if st.button("Process Text", key="process_text"):
        if text_input.strip():
            input_data = text_input
            data_type = "text"
        else:
            st.error("Please enter some text content")

with tab3:
    st.header("📁 File Upload")
    st.write("Upload various document types: PDF, DOC, DOCX, TXT")
    
    uploaded_file = st.file_uploader(
        "Choose a file", 
        type=["pdf", "doc", "docx", "txt"],
        key="file_uploader"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        st.write(f"📊 File size: {uploaded_file.size} bytes")
        st.write(f"📄 File type: {uploaded_file.type}")
        
        if st.button("🚀 Process File", key="process_file"):
            input_data = uploaded_file
            data_type = "file"
            start_processing = True

status_container = st.container()

faiss_directory = "faiss_store"
language_model = ChatOpenAI(
    api_key=openai_api_key, 
    temperature=0.7, 
    max_tokens=1000,
    model="gpt-3.5-turbo"
)
llm_embeddings = OpenAIEmbeddings(api_key=openai_api_key)

expected_minimum_chunks = 10

def split_text(loaded_data):
    primary_delimiters = ['\n\n', '\n', '.', ',']
    fallback_delimiters = [';', ' ', '|']

    doc_splitter = RecursiveCharacterTextSplitter(separators=primary_delimiters, chunk_size=1000, chunk_overlap=200)
    split_documents = doc_splitter.split_documents(loaded_data)
    if not split_documents or len(split_documents) < expected_minimum_chunks:
        doc_splitter = RecursiveCharacterTextSplitter(separators=fallback_delimiters, chunk_size=1000, chunk_overlap=200)
        split_documents = doc_splitter.split_documents(loaded_data)

    return split_documents

def process_pdf_file(uploaded_file):
    try:
        pdf_content = uploaded_file.read()

        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        text_content = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_content += page.extract_text() + "\n"
        
        if not text_content.strip():
            return None, "No text content found in PDF"
        
        from langchain.schema import Document
        document = Document(
            page_content=text_content,
            metadata={"source": uploaded_file.name, "type": "pdf"}
        )
        
        return [document], None
        
    except Exception as e:
        return None, f"Error processing PDF: {str(e)}"

def process_text_input(text_content):
    try:
        if not text_content.strip():
            return None, "Please enter some text"
        
        from langchain.schema import Document
        document = Document(
            page_content=text_content,
            metadata={"source": "Direct Text Input", "type": "text"}
        )
        
        return [document], None
        
    except Exception as e:
        return None, f"Error processing text: {str(e)}"

def enhance_documents_with_metadata(documents, source_type, source_name, additional_metadata=None):
    enhanced_docs = []
    for i, doc in enumerate(documents):
        metadata = {
            "source": doc.metadata.get("source", source_name),
            "source_type": source_type,
            "date_added": datetime.now().isoformat(),
            "chunk_index": i,
            "total_chunks": len(documents),
        }
        
        # Add any custom metadata
        if additional_metadata:
            metadata.update(additional_metadata)
        
        enhanced_doc = Document(
            page_content=doc.page_content,
            metadata=metadata
        )
        enhanced_docs.append(enhanced_doc)
    
    return enhanced_docs

def process_file_upload(uploaded_file):
    try:
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return process_pdf_file(uploaded_file)
        
        elif file_extension == 'txt':
            # Process TXT file
            text_content = uploaded_file.read().decode('utf-8')
            if not text_content.strip():
                return None, "No text content found in TXT file"
            
            from langchain.schema import Document
            document = Document(
                page_content=text_content,
                metadata={"source": uploaded_file.name, "type": "txt"}
            )
            return [document], None
        
        elif file_extension in ['doc', 'docx']:
            # Process DOC/DOCX file
            file_content = uploaded_file.read()
            
            if file_extension == 'docx':
                # Use python-docx for DOCX files
                doc = DocxDocument(io.BytesIO(file_content))
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
            else:
                return None, "DOC files are not yet supported. Please convert to DOCX or PDF format."
            
            if not text_content.strip():
                return None, "No text content found in document"
            
            from langchain.schema import Document
            document = Document(
                page_content=text_content,
                metadata={"source": uploaded_file.name, "type": file_extension}
            )
            return [document], None
        
        else:
            return None, f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        return None, f"Error processing file: {str(e)}"

def extract_keywords_with_yake(text, max_keywords=10):
    try:
        kw_extractor = yake.KeywordExtractor(
            lan="en",          
            n=3,                
            dedupLim=0.7,       
            top=20,             
            features=None       
        )
        
        keywords = kw_extractor.extract_keywords(text)
        if debug_yake:
            st.write("🔍 Raw YAKE output:", keywords[:5])
        
        extracted_keywords = []
        for kw_tuple in keywords[:max_keywords]:
            if isinstance(kw_tuple, tuple) and len(kw_tuple) >= 2:
                keyword = str(kw_tuple[1]).strip()
                if len(keyword) > 2 and keyword.replace(' ', '').isalpha():
                    extracted_keywords.append(keyword)
            elif isinstance(kw_tuple, str):
                keyword = kw_tuple.strip()
                if len(keyword) > 2 and keyword.replace(' ', '').isalpha():
                    extracted_keywords.append(keyword)
        
        return extracted_keywords
        
    except Exception as e:
        st.error(f"Error extracting keywords with YAKE: {str(e)}")
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'what', 'how', 'when', 'where', 'why', 'who', 'this', 'that', 'these', 'those'}
        filtered_words = [word for word in words if word not in stop_words]
        return filtered_words[:max_keywords]

def extract_keywords_simple(text, max_keywords=10):
    try:
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did',
            'will', 'would', 'could', 'should', 'may', 'might', 'can', 'what', 'how', 'when',
            'where', 'why', 'who', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she',
            'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'her',
            'its', 'our', 'their', 'from', 'into', 'during', 'including', 'until', 'against',
            'among', 'throughout', 'despite', 'towards', 'upon', 'concerning', 'to', 'of', 'in',
            'for', 'on', 'with', 'as', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
            'before', 'after', 'above', 'below', 'up', 'down', 'out', 'off', 'over', 'under',
            'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how',
            'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no',
            'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can',
            'will', 'just', 'don', 'should', 'now'
        }
        filtered_words = [word for word in words if word not in stop_words]
        word_counts = Counter(filtered_words)
        return [word for word, count in word_counts.most_common(max_keywords)]
        
    except Exception as e:
        st.error(f"Error in simple keyword extraction: {str(e)}")
        return []

def extract_keywords_from_query(query):
    yake_keywords = extract_keywords_with_yake(query, max_keywords=8)
    
    valid_yake_keywords = []
    for kw in yake_keywords:
        if isinstance(kw, str) and kw.replace(' ', '').isalpha() and len(kw) > 2:
            valid_yake_keywords.append(kw)
    
    if valid_yake_keywords:
        return valid_yake_keywords
    else:
        return extract_keywords_simple(query, max_keywords=8)

def extract_keywords_from_documents(documents):
    all_text = ""
    for doc in documents:
        all_text += doc.page_content + " "
    
    yake_keywords = extract_keywords_with_yake(all_text, max_keywords=10)
    
    valid_yake_keywords = []
    for kw in yake_keywords:
        if isinstance(kw, str) and kw.replace(' ', '').isalpha() and len(kw) > 2:
            valid_yake_keywords.append(kw)
    
    if valid_yake_keywords:
        return valid_yake_keywords
    else:
        return extract_keywords_simple(all_text, max_keywords=10)

@st.cache_resource
def load_cross_encoder():
    try:
        model = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
        return model
    except Exception as e:
        st.warning(f"Could not load cross-encoder model: {e}")
        return None

def bm25_rerank(query, documents, top_k=4):
    try:
        st.info("🔍 **BM25 Re-ranking**: Starting keyword-based ranking...")
        
        # Tokenize documents
        st.write("📝 Step 1: Tokenizing documents...")
        tokenized_docs = [doc.page_content.lower().split() for doc in documents]
        st.write(f"   ✅ Tokenized {len(tokenized_docs)} documents")
        st.write(f"   📊 Average tokens per document: {sum(len(d) for d in tokenized_docs) // len(tokenized_docs)}")
        
        # Create BM25 index
        st.write("📚 Step 2: Building BM25 index...")
        bm25 = BM25Okapi(tokenized_docs)
        st.write(f"   ✅ BM25 index created (k1=1.2, b=0.75)")
        
        # Tokenize query
        st.write("🔎 Step 3: Tokenizing query...")
        tokenized_query = query.lower().split()
        st.write(f"   ✅ Query tokens: {tokenized_query}")
        st.write(f"   📊 Number of query terms: {len(tokenized_query)}")
        
        # Get BM25 scores
        st.write("🎯 Step 4: Calculating BM25 scores...")
        scores = bm25.get_scores(tokenized_query)
        st.write(f"   ✅ Scores calculated for all documents")
        st.write(f"   📊 Score range: [{min(scores):.4f}, {max(scores):.4f}]")
        st.write(f"   📊 Mean score: {np.mean(scores):.4f}")
        
        # Sort documents by score
        st.write("📈 Step 5: Ranking documents by score...")
        ranked_indices = np.argsort(scores)[::-1][:top_k]
        ranked_docs = [(documents[i], float(scores[i])) for i in ranked_indices]
        
        st.success(f"✅ BM25 Re-ranking complete! Top {top_k} documents selected")
        st.write("🏆 Top 3 scores:")
        for i, (doc, score) in enumerate(ranked_docs[:3], 1):
            preview = doc.page_content[:100].replace('\n', ' ')
            st.write(f"   {i}. Score: {score:.4f} | Preview: {preview}...")
        
        return ranked_docs
    except Exception as e:
        st.error(f"❌ Error in BM25 re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def semantic_rerank(query, documents, embeddings_model, top_k=4):
    try:
        from sklearn.metrics.pairwise import cosine_similarity
        
        st.info("🧠 **Semantic Re-ranking**: Starting vector-based similarity matching...")
        
        # Get query embedding
        st.write("🔤 Step 1: Generating query embedding...")
        query_embedding = embeddings_model.embed_query(query)
        st.write(f"   ✅ Query embedded into {len(query_embedding)}-dimensional vector")
        st.write(f"   📊 Vector norm: {np.linalg.norm(query_embedding):.4f}")
        
        st.write("📄 Step 2: Generating document embeddings...")
        doc_embeddings = []
        for i, doc in enumerate(documents):
            emb = embeddings_model.embed_query(doc.page_content)
            doc_embeddings.append(emb)
            if i < 3:
                st.write(f"   ✅ Document {i+1} embedded (norm: {np.linalg.norm(emb):.4f})")
        st.write(f"   ✅ All {len(doc_embeddings)} documents embedded")
        
        st.write("📐 Step 3: Calculating cosine similarities...")
        similarities = [cosine_similarity([query_embedding], [doc_emb])[0][0] for doc_emb in doc_embeddings]
        st.write(f"   ✅ Similarities calculated")
        st.write(f"   📊 Similarity range: [{min(similarities):.4f}, {max(similarities):.4f}]")
        st.write(f"   📊 Mean similarity: {np.mean(similarities):.4f}")
        
        st.write("📈 Step 4: Ranking documents by similarity...")
        ranked_indices = np.argsort(similarities)[::-1][:top_k]
        ranked_docs = [(documents[i], float(similarities[i])) for i in ranked_indices]
        
        st.success(f"✅ Semantic Re-ranking complete! Top {top_k} documents selected")
        st.write("🏆 Top 3 similarities:")
        for i, (doc, score) in enumerate(ranked_docs[:3], 1):
            preview = doc.page_content[:100].replace('\n', ' ')
            st.write(f"   {i}. Similarity: {score:.4f} | Preview: {preview}...")
        
        return ranked_docs
    except Exception as e:
        st.error(f"❌ Error in semantic re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def cross_encoder_rerank(query, documents, cross_encoder, top_k=4):
    try:
        st.info("🤖 **Cross-Encoder Re-ranking**: Starting neural bi-encoder ranking...")
        
        if cross_encoder is None:
            st.warning("⚠️ Cross-encoder model not loaded, using fallback")
            return [(doc, 0.0) for doc in documents[:top_k]]
        st.write("🔗 Step 1: Creating query-document pairs...")
        pairs = [[query, doc.page_content] for doc in documents]
        st.write(f"   ✅ Created {len(pairs)} pairs for scoring")
        st.write(f"   📊 Average document length: {sum(len(doc.page_content) for doc in documents) // len(documents)} chars")
        st.write("🧮 Step 2: Computing cross-encoder scores...")
        st.write(f"   🔄 Processing {len(pairs)} pairs through transformer...")
        scores = cross_encoder.predict(pairs)
        st.write(f"   ✅ Scores computed!")
        st.write(f"   📊 Score range: [{min(scores):.4f}, {max(scores):.4f}]")
        st.write(f"   📊 Mean score: {np.mean(scores):.4f}")
        st.write(f"   📊 Std deviation: {np.std(scores):.4f}")
        
        st.write("📈 Step 3: Ranking documents by neural scores...")
        ranked_indices = np.argsort(scores)[::-1][:top_k]
        ranked_docs = [(documents[i], float(scores[i])) for i in ranked_indices]
        
        st.success(f"✅ Cross-Encoder Re-ranking complete! Top {top_k} documents selected")
        st.write("🏆 Top 3 neural scores:")
        for i, (doc, score) in enumerate(ranked_docs[:3], 1):
            preview = doc.page_content[:100].replace('\n', ' ')
            st.write(f"   {i}. Score: {score:.4f} | Preview: {preview}...")
        
        return ranked_docs
    except Exception as e:
        st.error(f"❌ Error in cross-encoder re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def hybrid_rerank(query, documents, embeddings_model, cross_encoder, top_k=4):
    try:
        st.info("🎯 **Hybrid Re-ranking**: Combining 3 methods with weighted ensemble...")
        st.write("📊 Step 1: Running all three ranking methods...")
        bm25_results = bm25_rerank(query, documents, top_k=len(documents))
        semantic_results = semantic_rerank(query, documents, embeddings_model, top_k=len(documents))
        cross_encoder_results = cross_encoder_rerank(query, documents, cross_encoder, top_k=len(documents))
        
        st.write("📏 Step 2: Normalizing scores to [0, 1] range...")
        def normalize_scores(results):
            scores = [score for _, score in results]
            if max(scores) > min(scores):
                normalized = [(doc, (score - min(scores)) / (max(scores) - min(scores))) 
                             for doc, score in results]
            else:
                normalized = [(doc, 0.5) for doc, _ in results]
            return normalized
        
        bm25_norm = normalize_scores(bm25_results)
        semantic_norm = normalize_scores(semantic_results)
        cross_encoder_norm = normalize_scores(cross_encoder_results)
        st.write(f"   ✅ All scores normalized")
        st.write("⚖️ Step 3: Applying weighted combination...")
        st.write(f"   📊 Weights: BM25=0.2, Semantic=0.3, Cross-Encoder=0.5")
        combined_scores = {}
        for i, doc in enumerate(documents):
            bm25_score = bm25_norm[i][1]
            semantic_score = semantic_norm[i][1]
            cross_score = cross_encoder_norm[i][1]
            
            combined_score = (0.2 * bm25_score) + (0.3 * semantic_score) + (0.5 * cross_score)
            combined_scores[i] = combined_score
            
            if i < 3:
                st.write(f"   Doc {i+1}: BM25={bm25_score:.3f}, Sem={semantic_score:.3f}, Cross={cross_score:.3f} → Final={combined_score:.3f}")
        st.write("📈 Step 4: Final ranking by combined scores...")
        ranked_indices = sorted(combined_scores.keys(), key=lambda x: combined_scores[x], reverse=True)[:top_k]
        ranked_docs = [(documents[i], float(combined_scores[i])) for i in ranked_indices]
        
        st.success(f"✅ Hybrid Re-ranking complete! Top {top_k} documents selected")
        st.write("🏆 Top 3 combined scores:")
        for i, (doc, score) in enumerate(ranked_docs[:3], 1):
            preview = doc.page_content[:100].replace('\n', ' ')
            st.write(f"   {i}. Score: {score:.4f} | Preview: {preview}...")
        
        return ranked_docs
    except Exception as e:
        st.error(f"❌ Error in hybrid re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def cohere_rerank(query, documents, cohere_api_key, top_k=4):
    try:
        if not cohere_api_key:
            st.warning("Cohere API key not provided. Using fallback method.")
            return hybrid_rerank(query, documents, None, None, top_k)
        
        co = cohere.Client(cohere_api_key)
        doc_texts = [doc.page_content for doc in documents]
        results = co.rerank(
            model="rerank-english-v3.0",
            query=query,
            documents=doc_texts,
            top_n=top_k,
            return_documents=True
        )
        
        ranked_docs = []
        for result in results.results:
            original_doc = documents[result.index]
            score = result.relevance_score
            ranked_docs.append((original_doc, float(score)))
        
        return ranked_docs
    except Exception as e:
        st.error(f"Error in Cohere re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def rankgpt_rerank(query, documents, llm_model, top_k=4):
    try:
        doc_list = []
        for i, doc in enumerate(documents):
            preview = doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content
            doc_list.append(f"[{i}] {preview}")
        
        prompt = f"""You are an expert at ranking document relevance for search queries.

Query: {query}

Documents to rank:
{chr(10).join(doc_list)}

Task: Rank these documents by relevance to the query. Return ONLY the indices in order from most to least relevant, separated by commas.
Example output: 2,0,4,1,3

Your ranking (indices only):"""
        
        response = llm_model.predict(prompt)
        try:
            ranked_indices = [int(x.strip()) for x in response.strip().split(',') if x.strip().isdigit()]
            ranked_indices = ranked_indices[:top_k]
        except:
            ranked_indices = list(range(min(top_k, len(documents))))
        
        ranked_docs = []
        for rank, idx in enumerate(ranked_indices):
            if idx < len(documents):
                score = 1.0 - (rank * 0.1)  # Decreasing score
                ranked_docs.append((documents[idx], score))
        
        return ranked_docs
    except Exception as e:
        st.error(f"Error in RankGPT re-ranking: {e}")
        return [(doc, 0.0) for doc in documents[:top_k]]

def reciprocal_rank_fusion(rankings_list, top_k=4, k=60):
    try:
        st.info("🔄 **Reciprocal Rank Fusion**: Combining multiple rankings...")
        st.write(f"📊 Number of ranking systems to fuse: {len(rankings_list)}")
        st.write(f"⚙️ RRF parameter k={k}")
        st.write("📚 Step 1: Identifying unique documents across all rankings...")
        all_docs = set()
        for ranking in rankings_list:
            for doc, _ in ranking:
                all_docs.add(id(doc))
        st.write(f"   ✅ Found {len(all_docs)} unique documents")
        
        # Calculate RRF scores
        st.write("🧮 Step 2: Calculating RRF scores...")
        rrf_scores = {}
        doc_map = {}
        
        for method_idx, ranking in enumerate(rankings_list):
            st.write(f"   📊 Processing ranking system {method_idx + 1}/{len(rankings_list)}...")
            for rank, (doc, score) in enumerate(ranking):
                doc_id = id(doc)
                doc_map[doc_id] = doc
                
                if doc_id not in rrf_scores:
                    rrf_scores[doc_id] = 0
                
                # RRF formula: 1/(k + rank)
                rrf_contribution = 1.0 / (k + rank + 1)
                rrf_scores[doc_id] += rrf_contribution
                
                if rank < 2:  # Show first 2 docs
                    st.write(f"      Doc at rank {rank+1}: RRF contribution = 1/({k}+{rank+1}) = {rrf_contribution:.6f}")
        
        st.write(f"   ✅ RRF scores calculated for all documents")
        st.write(f"   📊 Score range: [{min(rrf_scores.values()):.6f}, {max(rrf_scores.values()):.6f}]")
        

        st.write("📈 Step 3: Sorting by RRF scores...")
        sorted_docs = sorted(rrf_scores.items(), key=lambda x: x[1], reverse=True)[:top_k]
        
        ranked_docs = [(doc_map[doc_id], score) for doc_id, score in sorted_docs]
        
        st.success(f"✅ RRF Fusion complete! Top {top_k} documents selected")
        st.write("🏆 Top 3 RRF scores:")
        for i, (doc, score) in enumerate(ranked_docs[:3], 1):
            preview = doc.page_content[:100].replace('\n', ' ')
            st.write(f"   {i}. RRF Score: {score:.6f} | Preview: {preview}...")
        
        return ranked_docs
    except Exception as e:
        st.error(f"❌ Error in RRF: {e}")
        return rankings_list[0][:top_k] if rankings_list else []

def ultimate_hybrid_rerank(query, documents, embeddings_model, cross_encoder, llm_model, cohere_key, top_k=4):
    try:
        st.info("🏆 **Ultimate Hybrid**: Running up to 5 different re-ranking methods...")
        st.write(f"📊 **Methods Status Check:**")
        st.write(f"   • BM25: ✅ Always enabled")
        st.write(f"   • Semantic: ✅ Always enabled")
        st.write(f"   • Cross-Encoder: {'✅ Enabled' if cross_encoder else '❌ Disabled (model not loaded)'}")
        st.write(f"   • Cohere Re-rank: {'✅ Enabled' if cohere_key else '❌ Disabled (no API key)'}")
        st.write(f"   • RankGPT: ✅ Always enabled")
        st.write("")

        rankings = []
        methods_used = []
        
        # 1. BM25
        with st.spinner("🔄 Running BM25..."):
            bm25_ranking = bm25_rerank(query, documents, top_k=len(documents))
            rankings.append(bm25_ranking)
            methods_used.append("BM25")
        
        # 2. Semantic Similarity
        with st.spinner("🔄 Running Semantic Similarity..."):
            semantic_ranking = semantic_rerank(query, documents, embeddings_model, top_k=len(documents))
            rankings.append(semantic_ranking)
            methods_used.append("Semantic")
        
        # 3. Cross-Encoder
        if cross_encoder:
            with st.spinner("🔄 Running Cross-Encoder..."):
                cross_ranking = cross_encoder_rerank(query, documents, cross_encoder, top_k=len(documents))
                rankings.append(cross_ranking)
                methods_used.append("Cross-Encoder")
        else:
            st.warning("⚠️ Cross-Encoder skipped (model not loaded)")
        
        # 4. Cohere 
        if cohere_key:
            with st.spinner("🔄 Running Cohere Re-rank..."):
                cohere_ranking = cohere_rerank(query, documents, cohere_key, top_k=len(documents))
                rankings.append(cohere_ranking)
                methods_used.append("Cohere")
        else:
            st.warning("⚠️ Cohere Re-rank skipped (no API key in .env)")
        
        # 5. RankGPT
        with st.spinner("🔄 Running RankGPT..."):
            rankgpt_ranking = rankgpt_rerank(query, documents, llm_model, top_k=len(documents))
            rankings.append(rankgpt_ranking)
            methods_used.append("RankGPT")
        
        # Apply Reciprocal Rank Fusion
        st.success(f"✅ Combined {len(rankings)} ranking methods using RRF: {', '.join(methods_used)}")
        final_ranking = reciprocal_rank_fusion(rankings, top_k=top_k)
        
        return final_ranking
    except Exception as e:
        st.error(f"Error in ultimate hybrid re-ranking: {e}")
        return hybrid_rerank(query, documents, embeddings_model, cross_encoder, top_k)

def calculate_confidence_score(ranked_docs, query, strategy):
    try:
        if not ranked_docs or len(ranked_docs) < 2:
            return 0.0
        
        scores = [score for _, score in ranked_docs]
        if "RRF" in strategy or "Ultimate" in strategy:
            normalized_scores = [s * 15 for s in scores]  
            st.write(f"🔍 **RRF Score Normalization**: Original top={scores[0]:.6f}, Normalized={normalized_scores[0]:.4f}")
        elif "BM25" in strategy:
            max_expected = max(max(scores), 10.0)
            normalized_scores = [s / max_expected for s in scores]
        else:
            normalized_scores = scores
        top_score = min(normalized_scores[0], 1.0)
        st.write(f"   📊 Top Score Component: {top_score:.4f} (40% weight)")

        if len(normalized_scores) > 1:
            score_gap = normalized_scores[0] - normalized_scores[1]
            separation_score = min(score_gap * 5, 1.0)  # Larger gap = more confident
            st.write(f"   📊 Score Separation: {score_gap:.4f} → {separation_score:.4f} (30% weight)")
        else:
            separation_score = 0.5
        
        top_3_avg = np.mean(normalized_scores[:min(3, len(normalized_scores))])
        consistency_score = min(top_3_avg, 1.0)
        st.write(f"   📊 Top-3 Consistency: {consistency_score:.4f} (30% weight)")
        
        confidence = (0.4 * top_score) + (0.3 * separation_score) + (0.3 * consistency_score)
        
        st.write(f"   🎯 **Final Calculation**: 0.4×{top_score:.4f} + 0.3×{separation_score:.4f} + 0.3×{consistency_score:.4f}")
        st.write(f"   🎯 **Raw Confidence**: {confidence:.4f}")
        
        final_confidence = min(confidence * 100, 100.0)
        st.write(f"   ✅ **Confidence Score**: {final_confidence:.2f}%")
        
        return round(final_confidence, 2)
    except Exception as e:
        st.error(f"Error calculating confidence score: {e}")
        return 0.0

def fetch_news_articles(keywords, num_articles=5):
    if not keywords:
        return []
    
    try:
        search_query = " OR ".join(keywords)
        from_date = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')
        
        url = "https://newsapi.org/v2/everything"
        params = {
            'q': search_query,
            'from': from_date,
            'sortBy': 'publishedAt',
            'language': 'en',
            'pageSize': num_articles,
            'apiKey': news_api_key
        }
        
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        
        data = response.json()
        
        if data['status'] == 'ok' and data['articles']:
            return data['articles']
        else:
            return []
            
    except Exception as e:
        st.error(f"❌ Error fetching news: {str(e)}")
        return []

def display_news_articles(articles):
    if not articles:
        return
    
    st.subheader("📰 Related News Articles")
    
    for i, article in enumerate(articles, 1):
        with st.expander(f"📄 {article['title'][:80]}..." if len(article['title']) > 80 else f"📄 {article['title']}"):
            st.write(f"**Source:** {article['source']['name']}")
            st.write(f"**Published:** {article['publishedAt'][:10]}")
            
            if article.get('description'):
                st.write(f"**Description:** {article['description']}")
            
            if article.get('url'):
                st.write(f"**Read more:** [Link]({article['url']})")
            
            if article.get('urlToImage'):
                st.image(article['urlToImage'], width=300)
if input_data is not None and data_type is not None:
    status_container.info(f"Processing {data_type}...")
    
    if show_processing_steps:
        st.info("🔍 **Debug: Starting Processing Pipeline**")
    
    try:
        loaded_data = None
        
        if data_type == "urls":
            if show_processing_steps:
                st.info(f"📥 **Step 1: Loading URLs** - Processing {len(input_data)} URL(s)")
                for i, url in enumerate(input_data, 1):
                    st.write(f"   {i}. {url}")
            
            url_loader = WebBaseLoader(input_data)
        loaded_data = url_loader.load()

        if show_processing_steps:
            st.info(f"✅ **Step 1 Complete** - Loaded {len(loaded_data)} raw documents")
            
        elif data_type == "pdf":
            if show_processing_steps:
                st.info(f"📄 **Step 1: Processing PDF** - {input_data.name}")
                st.write(f"   File size: {input_data.size} bytes")
            
            loaded_data, error = process_pdf_file(input_data)
            if error:
                status_container.error(f"❌ {error}")
                loaded_data = None
            else:
                if show_processing_steps:
                    st.info(f"✅ **Step 1 Complete** - Extracted text from PDF")
                
        elif data_type == "text":
            if show_processing_steps:
                st.info(f"📝 **Step 1: Processing Text** - {len(input_data)} characters")
                st.write(f"   Preview: {input_data[:100]}...")
            
            loaded_data, error = process_text_input(input_data)
            if error:
                status_container.error(f"❌ {error}")
                loaded_data = None
            else:
                if show_processing_steps:
                    st.info(f"✅ **Step 1 Complete** - Created document from text")
        
        elif data_type == "file":
            if show_processing_steps:
                st.info(f"📁 **Step 1: Processing File** - {input_data.name}")
                st.write(f"   File size: {input_data.size} bytes")
                st.write(f"   File type: {input_data.type}")
            
            loaded_data, error = process_file_upload(input_data)
            if error:
                status_container.error(f"❌ {error}")
                loaded_data = None
            else:
                if show_processing_steps:
                    st.info(f"✅ **Step 1 Complete** - Processed file: {input_data.name}")
        
        if loaded_data:
            if show_processing_steps:
                st.info(f"✂️ **Step 2: Text Splitting** - Creating chunks from {len(loaded_data)} document(s)")
        
        split_documents = split_text(loaded_data)

        if not split_documents:
            status_container.error("Failed to split documents into chunks")
        else:
            if show_processing_steps:
                st.info(f"✅ **Step 2 Complete** - Created {len(split_documents)} chunks")
                st.write(f"   Average chunk size: {sum(len(chunk.page_content) for chunk in split_documents) // len(split_documents)} characters")
            
            if show_storage_info:
                st.info(f"💾 **Step 3: Creating Vector Store** - Generating embeddings and saving to FAISS")
            
            vectorindex_openai = FAISS.from_documents(split_documents, llm_embeddings)
            
            if show_storage_info:
                st.info(f"💾 **Step 4: Saving to Disk** - Storing in '{faiss_directory}' directory")
                
        vectorindex_openai.save_local("faiss_store")
        
        if show_storage_info:
            st.info(f"📁 **Storage Location:** `{os.path.abspath(faiss_directory)}`")
            st.write(f"   Files created:")
            st.write(f"   - `{faiss_directory}/index.faiss` (vector index)")
            st.write(f"   - `{faiss_directory}/index.pkl` (metadata)")
        
        source_info = ""
        if data_type == "urls":
            source_info = f"from {len(input_data)} URL(s)"
        elif data_type == "pdf":
            source_info = f"from PDF: {input_data.name}"
        elif data_type == "text":
            source_info = "from direct text input"
        elif data_type == "file":
            source_info = f"from file: {input_data.name}"
        
        status_container.success(f"✅ {data_type.title()} processed successfully! Created {len(split_documents)} document chunks {source_info}.")
        
        if enable_debug:
            st.success("🎉 **Processing Complete!** Ready to answer questions.")
                
    except Exception as e:
        status_container.error(f"❌ Error processing {data_type}: {str(e)}")
        if data_type == "urls":
            st.error("Make sure the URLs are accessible and contain readable content.")
        elif data_type == "pdf":
            st.error("Make sure the PDF file is valid and contains readable text.")
        elif data_type == "text":
            st.error("There was an error processing your text input.")
        elif data_type == "file":
            st.error("There was an error processing your file. Make sure the file is valid and contains readable content.")

if os.path.isdir(faiss_directory):
    
    try:
        sample_index = FAISS.load_local(faiss_directory, llm_embeddings)
        
    except:
        pass

st.divider()
st.header("💬 Ask Questions")

user_query = st.text_input("🔍 Type your query here", placeholder="Ask anything about your processed content...")

if user_query:
    if os.path.isdir(faiss_directory):
        try:
            is_followup = st.session_state.conversation_memory.detect_followup_query(user_query)
            if is_followup:
                st.info("🔗 Detected follow-up question! Using conversation context...")
            

            adaptive_k = st.session_state.adaptive_context.get_context_window_size(user_query, base_k=num_candidates)
            if adaptive_k != num_candidates:
                st.write(f"📊 Adaptive retrieval: Adjusted from {num_candidates} to {adaptive_k} documents")

            enhanced_query = st.session_state.adaptive_context.build_enhanced_query(user_query)
            if enhanced_query != user_query:
                with st.expander("👁️ View Enhanced Query with Context"):
                    st.text(enhanced_query)

            if show_processing_steps:
                st.info(f"🔍 **Query Processing: Loading Vector Store**")
                st.write(f"   Loading from: `{os.path.abspath(faiss_directory)}`")
            
            loaded_faiss_index = FAISS.load_local(faiss_directory, llm_embeddings)

            if st.session_state.metadata_filter.filters:
                st.info(f"🔎 Applying {len(st.session_state.metadata_filter.filters)} metadata filter(s)")

                filter_candidates = loaded_faiss_index.similarity_search(enhanced_query, k=adaptive_k*3)

                initial_docs = [
                    doc for doc in filter_candidates 
                    if st.session_state.metadata_filter.matches(doc.metadata)
                ][:adaptive_k]
                
                st.write(f"   Filtered from {len(filter_candidates)} to {len(initial_docs)} documents")
                
                if not initial_docs:
                    st.warning("⚠️ No documents match your filters! Try relaxing the filters.")
                    st.stop()
            else:
                if show_processing_steps:
                    st.info(f"🔍 **Query Processing: Initial Retrieval**")
                    st.write(f"   Retrieving top {adaptive_k} candidate chunks")

                    initial_docs = loaded_faiss_index.similarity_search(enhanced_query, k=adaptive_k)

                if enable_reranking:
                    if show_processing_steps:
                        st.info(f"🎯 **Re-Ranking: Applying {reranking_strategy}**")
                        
 
                        cross_encoder = load_cross_encoder() if ("Cross-Encoder" in reranking_strategy or "Hybrid" in reranking_strategy or "Ultimate" in reranking_strategy) else None
                        if reranking_strategy == "🏆 Ultimate Hybrid (All Methods + RRF)":
                            ranked_docs_with_scores = ultimate_hybrid_rerank(user_query, initial_docs, llm_embeddings, cross_encoder, language_model, cohere_api_key, top_k=num_final_results)
                        elif reranking_strategy == "Hybrid (BM25 + Semantic + Cross-Encoder)":
                            ranked_docs_with_scores = hybrid_rerank(user_query, initial_docs, llm_embeddings, cross_encoder, top_k=num_final_results)
                        elif reranking_strategy == "Cohere Re-rank (Industry-leading)":
                            ranked_docs_with_scores = cohere_rerank(user_query, initial_docs, cohere_api_key, top_k=num_final_results)
                        elif reranking_strategy == "RankGPT (LLM-based)":
                            ranked_docs_with_scores = rankgpt_rerank(user_query, initial_docs, language_model, top_k=num_final_results)
                        elif reranking_strategy == "BM25 Only":
                            ranked_docs_with_scores = bm25_rerank(user_query, initial_docs, top_k=num_final_results)
                        elif reranking_strategy == "Semantic Only":
                            ranked_docs_with_scores = semantic_rerank(user_query, initial_docs, llm_embeddings, top_k=num_final_results)
                        elif reranking_strategy == "Cross-Encoder Only":
                            ranked_docs_with_scores = cross_encoder_rerank(user_query, initial_docs, cross_encoder, top_k=num_final_results)
                        
  
                        final_docs = [doc for doc, score in ranked_docs_with_scores]

                        confidence_score = calculate_confidence_score(ranked_docs_with_scores, user_query, reranking_strategy)
                        
                        if show_processing_steps:
                            st.info(f"🎯 **Re-Ranking: Complete**")
                            st.write(f"   Final documents: {len(final_docs)}")
                            st.write(f"   Confidence score: {confidence_score}%")
                        else:
                            final_docs = initial_docs[:num_final_results]
                            ranked_docs_with_scores = [(doc, 0.0) for doc in final_docs]
                            confidence_score = 0.0
                        
                        if show_chunks:
                            st.info(f"📄 **Retrieved Chunks (Top {len(final_docs)}):**")
                            for i, (doc, score) in enumerate(ranked_docs_with_scores, 1):
                                score_display = f"{score:.4f}" if enable_reranking else "N/A"
                                with st.expander(f"Chunk {i} (Score: {score_display})"):
                                    st.write(f"**Source:** {doc.metadata.get('source', 'Unknown')}")
                                    st.write(f"**Content Length:** {len(doc.page_content)} characters")
                                    if enable_reranking:
                                        st.write(f"**Relevance Score:** {score:.4f}")
                                    st.write(f"**Content Preview:**")
                                    st.write(doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content)
                        
                        if show_processing_steps:
                            st.info(f"🔍 **Query Processing: Generating Answer**")
                        
                        from langchain.schema import Document
                        from langchain.callbacks.base import BaseCallbackHandler
                        
                        qa_chain = RetrievalQA.from_chain_type(
                            llm=language_model, 
                            chain_type="stuff", 
                            retriever=loaded_faiss_index.as_retriever(search_kwargs={"k": num_final_results}),
                            return_source_documents=True
                        )
                        with st.spinner("Generating answer based on re-ranked documents..."):
                            from langchain.prompts import PromptTemplate

                            context = "\n\n".join([f"Document {i+1}:\n{doc.page_content}" for i, doc in enumerate(final_docs)])
                            prompt = f"""Based on the following context, please answer the question.
                            
Context:
{context}

Question: {user_query}

Answer:"""
                            
                            answer = language_model.predict(prompt)
                            query_result = {"result": answer, "source_documents": final_docs}
                        
                        st.subheader("🤖 Answer")

                        if enable_reranking and show_confidence_scores:
                            if isinstance(confidence_score, (int, float)) and not (confidence_score != confidence_score):  # NaN check
                                confidence_score = max(0.0, min(100.0, confidence_score))  # Clamp to 0-100
                                
                                if confidence_score >= 75:
                                    confidence_color = "🟢"
                                    confidence_label = "High"
                                elif confidence_score >= 50:
                                    confidence_color = "🟡"
                                    confidence_label = "Medium"
                                else:
                                    confidence_color = "🔴"
                                    confidence_label = "Low"
                                
                                st.markdown(f"**{confidence_color} Confidence Score: {confidence_score:.1f}% ({confidence_label})**")
                                st.progress(min(confidence_score / 100, 1.0))
                                st.write("")
                            else:
                                st.info("⚠️ Confidence score unavailable for this result")
                        
                        st.write(query_result["result"])

                        st.session_state.continuity_system.process_message(
                            user_query, 
                            role="user",
                            message_type="rag_question"
                        )
                        st.session_state.continuity_system.process_message(
                            query_result["result"][:500],  
                            role="assistant",
                            message_type="rag_answer"
                        )
                        st.success("📌 Question tracked in Conversation Memory!")
                        st.divider()
                        st.subheader("📝 Was this answer helpful?")
                        
                        feedback_col1, feedback_col2, feedback_col3 = st.columns([1, 1, 2])
                        
                        with feedback_col1:
                            if st.button("👍 Yes", key="thumbs_up"):
                                st.session_state.feedback_storage.add_feedback(
                                    query=user_query,
                                    answer=query_result["result"],
                                    thumbs='up',
                                    retrieved_docs=[doc.metadata.get('source') for doc in final_docs],
                                    confidence_score=confidence_score if enable_reranking else None,
                                    reranking_strategy=reranking_strategy if enable_reranking else "None",
                                    session_id=st.session_state.session_id
                                )
                            st.success("✅ Thank you for your feedback!")
                        
                        with feedback_col2:
                            if st.button("👎 No", key="thumbs_down"):
                                st.session_state.feedback_storage.add_feedback(
                                    query=user_query,
                                    answer=query_result["result"],
                                    thumbs='down',
                                    retrieved_docs=[doc.metadata.get('source') for doc in final_docs],
                                    confidence_score=confidence_score if enable_reranking else None,
                                    reranking_strategy=reranking_strategy if enable_reranking else "None",
                                    session_id=st.session_state.session_id
                                )
                            st.warning("📝 Feedback recorded. We'll work to improve!")
                        
                        with st.expander("💬 Provide detailed feedback (optional)"):
                            feedback_rating = st.slider("Rate this answer", 1, 5, 3, key="feedback_rating")
                            feedback_comment = st.text_area("Additional comments", key="feedback_comment", placeholder="What could be improved?")
                            
                            if st.button("Submit Detailed Feedback", key="submit_feedback"):
                                st.session_state.feedback_storage.add_feedback(
                                    query=user_query,
                                    answer=query_result["result"],
                                    rating=feedback_rating,
                                    comment=feedback_comment if feedback_comment else None,
                                    retrieved_docs=[doc.metadata.get('source') for doc in final_docs],
                                    confidence_score=confidence_score if enable_reranking else None,
                                    reranking_strategy=reranking_strategy if enable_reranking else "None",
                                    session_id=st.session_state.session_id
                                )
                            st.success("✅ Detailed feedback saved! Thank you!")

                        query_stats = st.session_state.feedback_storage.get_query_stats(user_query)
                        if query_stats and query_stats['total_queries'] > 1:
                            with st.expander(f"📊 Historical Performance ({query_stats['total_queries']} similar queries)"):
                                perf_col1, perf_col2, perf_col3 = st.columns(3)
                                with perf_col1:
                                    st.metric("👍 Positive", query_stats['positive_feedback'])
                                with perf_col2:
                                    st.metric("👎 Negative", query_stats['negative_feedback'])
                                with perf_col3:
                                    st.metric("⭐ Avg Rating", f"{query_stats['avg_rating']:.1f}/5")
                        
                        st.session_state.conversation_memory.add_turn(
                            query=user_query,
                            answer=query_result["result"],
                            retrieved_docs=[doc.metadata.get('source') for doc in final_docs],
                            confidence=confidence_score if enable_reranking else 0.0
                        )
                        
                        st.divider()
                        
                        if show_processing_steps:
                            st.info(f"🔑 **Keyword Extraction: Starting**")
                        
                        query_keywords = extract_keywords_from_query(user_query)
                        document_keywords = extract_keywords_from_documents(query_result.get("source_documents", []))
                        all_keywords = list(set(query_keywords + document_keywords))
                        
                        if show_processing_steps:
                            st.info(f"🔑 **Keyword Extraction: Complete**")
                            st.write(f"   Query keywords: {len(query_keywords)}")
                            st.write(f"   Document keywords: {len(document_keywords)}")
                            st.write(f"   Combined unique keywords: {len(all_keywords)}")
                        
                        if all_keywords and enable_news:
                            if show_keywords:
                                st.info(f"🔍 Query keywords: {', '.join(query_keywords[:5])}")
                                st.info(f"📄 Document keywords: {', '.join(document_keywords[:5])}")
                                st.info(f"🎯 Combined keywords: {', '.join(all_keywords[:8])}")
                                
                                if show_processing_steps:
                                    st.info(f"📰 **News Search: Starting**")
                                    st.write(f"   Searching for: {', '.join(all_keywords[:8])}")
                                    st.write(f"   Number of articles requested: {num_news_articles}")
                                
                                st.info(f"🔍 Searching for news related to: {', '.join(all_keywords[:8])}")
                                with st.spinner("Fetching related news articles..."):
                                    news_articles = fetch_news_articles(all_keywords, num_articles=num_news_articles)
                                    
                                    if show_processing_steps:
                                        st.info(f"📰 **News Search: Complete**")
                                        st.write(f"   Articles found: {len(news_articles)}")
                                    
                                    if news_articles:
                                        display_news_articles(news_articles)
                                    else:
                                        st.info("No recent news articles found for these keywords.")
                            if query_result.get("source_documents"):
                                st.subheader("📚 Sources")

                                unique_sources = {}
                                for i, doc in enumerate(query_result["source_documents"], 1):
                                    source_url = doc.metadata.get('source', 'Unknown')
                                    if source_url not in unique_sources:
                                        unique_sources[source_url] = {
                                            'chunks': [],
                                            'chunk_count': 0
                                        }
                                    unique_sources[source_url]['chunks'].append((i, doc))
                                    unique_sources[source_url]['chunk_count'] += 1
                                

                                st.write("**Referenced Sources:**")
                                for source_url, data in unique_sources.items():
                                    chunk_count = data['chunk_count']
                                    if chunk_count > 1:
                                        st.write(f"• **{source_url}** ({chunk_count} chunks)")
                                    else:
                                        st.write(f"• **{source_url}**")
                                for source_url, data in unique_sources.items():
                                    with st.expander(f"📄 {source_url}"):
                                        for chunk_num, doc in data['chunks']:
                                            st.write(f"**Chunk {chunk_num}:**")
                                            st.write(doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content)
                                            st.write("---")
        except Exception as e:
            st.error(f"❌ Error loading FAISS index: {str(e)}")
            st.info("Try reprocessing the URLs to recreate the index.")
        # else:
        #     status_container.error("FAISS index not found. Please process the URLs first.")


with app_tab2:
    st.header("📊 System Analytics & Performance")

    st.subheader("📈 Overall Statistics")
    recent_feedback = st.session_state.feedback_storage.get_recent_feedback(100)
    
    if recent_feedback:
        thumbs_up = sum(1 for f in recent_feedback if f['thumbs'] == 'up')
        thumbs_down = sum(1 for f in recent_feedback if f['thumbs'] == 'down')
        rated = [f for f in recent_feedback if f['rating'] is not None]
        avg_rating = sum(f['rating'] for f in rated) / len(rated) if rated else 0
        
        metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
        
        with metric_col1:
            st.metric("📊 Total Queries", len(recent_feedback))
        
        with metric_col2:
            st.metric("👍 Positive", thumbs_up)
        
        with metric_col3:
            st.metric("👎 Negative", thumbs_down)
        
        with metric_col4:
            if thumbs_up + thumbs_down > 0:
                success_rate = thumbs_up / (thumbs_up + thumbs_down) * 100
                st.metric("✅ Success Rate", f"{success_rate:.1f}%")
            else:
                st.metric("✅ Success Rate", "N/A")
        
        if rated:
            st.metric("⭐ Average Rating", f"{avg_rating:.2f}/5")

        st.subheader("🎯 Confidence Score Distribution")
        confidence_scores = [f['confidence'] for f in recent_feedback if f['confidence'] is not None]
        
        if confidence_scores:
            import pandas as pd
            df = pd.DataFrame({
                'Confidence Score': confidence_scores
            })
            st.bar_chart(df['Confidence Score'])
        else:
            st.info("No confidence score data available yet")
        
        st.subheader("📝 Recent Feedback (Last 10)")
        for i, feedback in enumerate(recent_feedback[:10], 1):
            with st.expander(f"{i}. {feedback['query'][:60]}... ({feedback['timestamp'][:10]})"):
                st.write(f"**Query:** {feedback['query']}")
                st.write(f"**Answer:** {feedback['answer'][:200]}...")
                
                feedback_status_col1, feedback_status_col2, feedback_status_col3 = st.columns(3)
                with feedback_status_col1:
                    if feedback['thumbs']:
                        st.write(f"**Feedback:** {'👍' if feedback['thumbs'] == 'up' else '👎'}")
                    
                    with feedback_status_col2:
                        if feedback['rating']:
                            st.write(f"**Rating:** {feedback['rating']}/5 ⭐")
                    
                    with feedback_status_col3:
                        if feedback['confidence']:
                            st.write(f"**Confidence:** {feedback['confidence']:.0f}%")
                    
                    if feedback['comment']:
                        st.write(f"**Comment:** {feedback['comment']}")
                    else:
                        st.info("📭 No feedback data yet. Start asking questions and providing feedback!")
    
    st.divider()
    
    st.subheader("⚠️ Queries Needing Improvement")
    st.caption("Queries with lowest success rates (minimum 3 occurrences)")
    
    worst_queries = st.session_state.feedback_storage.get_worst_performing_queries(10)
    
    if worst_queries:
        for i, query_data in enumerate(worst_queries, 1):
            success_rate = query_data['success_rate'] * 100
            
            if success_rate < 40:
                status_emoji = "🔴"
            elif success_rate < 70:
                status_emoji = "🟡"
            else:
                status_emoji = "🟢"
            
            with st.expander(f"{i}. {status_emoji} {query_data['query']} (Success: {success_rate:.0f}%)"):
                perf_col1, perf_col2, perf_col3, perf_col4 = st.columns(4)
                
                with perf_col1:
                    st.metric("Total", query_data['total_queries'])
                
                with perf_col2:
                    st.metric("👍 Positive", query_data['positive_feedback'])
                
                with perf_col3:
                    st.metric("👎 Negative", query_data['negative_feedback'])
                
                with perf_col4:
                    st.metric("⭐ Avg Rating", f"{query_data['avg_rating']:.1f}/5")
    else:
        st.info("Not enough data yet to identify problem queries")

with app_tab3:
    st.header("⚙️ Advanced Settings & Configuration")

    st.subheader("🔧 System Configuration")
    
    config_col1, config_col2 = st.columns(2)
    
    with config_col1:
        st.write("**API Keys:**")
        st.write(f"- OpenAI: {'✅ Configured' if openai_api_key else '❌ Missing'}")
        st.write(f"- News API: {'✅ Configured' if news_api_key else '❌ Missing'}")
        st.write(f"- Cohere: {'✅ Configured' if cohere_api_key else '❌ Missing'}")
    
    with config_col2:
        st.write("**Storage:**")
        if os.path.isdir(faiss_directory):
            st.write(f"- Vector Store: ✅ Active")
            st.write(f"- Location: `{os.path.abspath(faiss_directory)}`")
        else:
            st.write(f"- Vector Store: ❌ Not created")
        
        st.write(f"- Feedback DB: ✅ Active")
        st.write(f"- Session ID: `{st.session_state.session_id}`")
    
    st.divider()
    st.subheader("💬 Conversation Memory Settings")
    
    new_max_history = st.slider(
        "Max conversation turns to remember",
        min_value=5,
        max_value=20,
        value=st.session_state.conversation_memory.max_history,
        help="Number of previous conversation turns to keep in memory"
    )
    
    new_max_tokens = st.slider(
        "Max context tokens",
        min_value=500,
        max_value=4000,
        value=st.session_state.conversation_memory.max_tokens,
        help="Maximum number of tokens to use from conversation history"
    )
    
    if st.button("Apply Memory Settings"):
        st.session_state.conversation_memory.max_history = new_max_history
        st.session_state.conversation_memory.max_tokens = new_max_tokens
        st.success("✅ Settings applied!")
    
    st.divider()

    st.subheader("🗄️ Data Management")
    
    data_col1, data_col2 = st.columns(2)
    
    with data_col1:
        st.write("**Export Data:**")
        if st.button("📥 Export Feedback Data"):
            feedback_data = st.session_state.feedback_storage.get_recent_feedback(1000)
            if feedback_data:
                import json
                json_data = json.dumps(feedback_data, indent=2)
                st.download_button(
                    label="Download Feedback JSON",
                    data=json_data,
                    file_name=f"feedback_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
            else:
                st.info("No feedback data to export")
    
    with data_col2:
        st.write("**Clear Data:**")
        if st.button("🗑️ Clear Conversation History", type="secondary"):
            st.session_state.conversation_memory.clear()
            st.success("✅ Conversation history cleared!")
        
        if st.button("⚠️ Clear All Filters", type="secondary"):
            st.session_state.metadata_filter.clear()
            st.success("✅ All filters cleared!")
    
    st.divider()

    st.subheader("ℹ️ System Information")
    
    info_col1, info_col2 = st.columns(2)
    
    with info_col1:
        st.write("**Features Enabled:**")
        st.write(f"- Advanced Re-Ranking: {'✅' if enable_reranking else '❌'}")
        st.write(f"- News Integration: {'✅' if enable_news else '❌'}")
        st.write(f"- Metadata Filtering: {'✅' if st.session_state.metadata_filter.filters else '❌'}")
        st.write(f"- Conversation Memory: ✅")
        st.write(f"- Feedback Collection: ✅")
    
    with info_col2:
        st.write("**Performance:**")
        st.write(f"- Re-Ranking Strategy: {reranking_strategy}")
        st.write(f"- Candidates Retrieved: {num_candidates}")
        st.write(f"- Final Results: {num_final_results}")
        st.write(f"- Conversation History: {len(st.session_state.conversation_memory.history)} turns")
    
    st.divider()

with app_tab4:
    st.header("🔍 Metadata Extraction with Crawl4AI")
    
    st.markdown("""
    Extract comprehensive metadata from web documents using Crawl4AI.
    Enter a URL or paste document content to see all extracted metadata.
    """)
    
    extraction_method = st.radio("Choose input method:", ["🌐 URL Extraction", "📄 Document Text", "📋 Test Examples"])
    
    if extraction_method == "🌐 URL Extraction":
        st.subheader("🌐 Extract Metadata from URL")
        
        url_input = st.text_input("Enter URL:", placeholder="https://example.com/article")
        
        if url_input and st.button("🔍 Extract Metadata"):
            try:
                import asyncio
                
                st.info("🔄 Extracting metadata from URL...")
                

                extractor = Crawl4AIMetadataExtractor()
                metadata = asyncio.run(extractor.extract_metadata_from_url(url_input))
                
                st.success("✅ Metadata extracted successfully!")

                meta_col1, meta_col2 = st.columns(2)
                
                with meta_col1:
                    st.write("### 📋 Basic Information")
                    st.write(f"**Title:** {metadata.get('title', 'N/A')}")
                    st.write(f"**Source:** {metadata.get('source', 'N/A')}")
                    st.write(f"**Language:** {metadata.get('language', 'N/A')}")
                    st.write(f"**Content Length:** {metadata.get('content_length', 0):,} characters")
                    if metadata.get('author'):
                        st.write(f"**Author:** {metadata['author']}")
                    if metadata.get('publication_date'):
                        st.write(f"**Published:** {metadata['publication_date']}")
                
                with meta_col2:
                    st.write("### 📊 Content Metrics")
                    st.write(f"**Read Time:** {metadata.get('estimated_read_time_minutes', 0)} minutes")
                    st.write(f"**Has Code:** {'✅ Yes' if metadata.get('has_code_blocks') else '❌ No'}")
                    st.write(f"**Has Lists:** {'✅ Yes' if metadata.get('has_lists') else '❌ No'}")
                    st.write(f"**Has Headers:** {'✅ Yes' if metadata.get('has_headers') else '❌ No'}")
                    st.write(f"**Links Found:** {metadata.get('link_count', 0)}")
                    st.write(f"**Images Found:** {metadata.get('image_count', 0)}")

                st.write("---")
                st.write("### 🎯 Web Quality Assessment")
                
                quality_assessor = WebMetadataQualityAssessor()
                assessment = quality_assessor.assess(metadata)
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("Overall Score", f"{assessment['overall_score']:.2f}", delta=f"{assessment['overall_score']*100:.0f}%")
                
                with col2:
                    st.metric("Completeness", f"{assessment['completeness']:.2f}", delta=f"{assessment['completeness']*100:.0f}%")
                
                with col3:
                    st.metric("Accuracy", f"{assessment['accuracy']:.2f}", delta=f"{assessment['accuracy']*100:.0f}%")
                
                with col4:
                    st.metric("Enrichment", f"{assessment['enrichment_level']:.2f}", delta=f"{assessment['enrichment_level']*100:.0f}%")
                
                # Issues
                if assessment.get('issues'):
                    st.warning("### ⚠️ Issues Found:")
                    for issue in assessment['issues']:
                        st.write(f"  • {issue}")
                
                # Raw metadata display
                with st.expander("📝 Full Metadata JSON"):
                    import json
                    st.json(metadata)
                
                # Enriched metadata
                st.write("---")
                st.write("### 🛠️ Enriched Metadata")
                
                enricher = MetadataEnricher()
                enriched = enricher.enrich_metadata(metadata.get('content', ''), metadata)
                
                enrich_col1, enrich_col2 = st.columns(2)
                
                with enrich_col1:
                    st.write("**Text Analysis:**")
                    st.write(f"  • Word Count: {enriched.get('word_count', 0):,}")
                    st.write(f"  • Sentence Count: {enriched.get('sentence_count', 0)}")
                    st.write(f"  • Avg Word Length: {enriched.get('avg_word_length', 0):.1f}")
                    st.write(f"  • Avg Sentence Length: {enriched.get('avg_sentence_length', 0):.1f}")
                
                with enrich_col2:
                    st.write("**Content Quality:**")
                    st.write(f"  • Readability: {enriched.get('readability_score', 0):.1f}/100")
                    st.write(f"  • Sentiment: {enriched.get('sentiment', 'N/A')}")
                    st.write(f"  • Language: {enriched.get('language', 'N/A')}")
                    if enriched.get('key_phrases'):
                        st.write(f"  • Key Phrases: {', '.join(enriched['key_phrases'][:5])}")
                
            except Exception as e:
                st.error(f"❌ Error extracting metadata: {str(e)}")
                st.info("💡 Tip: Make sure Crawl4AI is installed: `pip install crawl4ai`")
    
    elif extraction_method == "📄 Document Text":
        st.subheader("📄 Analyze Document Metadata")
        
        doc_text = st.text_area(
            "Paste document content:",
            height=300,
            placeholder="Paste your document text here..."
        )
        
        if doc_text and st.button("📊 Analyze"):
            st.write("### 📊 Text Analysis")
            
            enricher = MetadataEnricher()
            test_metadata = {
                'source': 'Manual Input',
                'title': 'Analyzed Document',
                'source_type': 'text'
            }
            
            enriched = enricher.enrich_metadata(doc_text, test_metadata)

            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Word Count", f"{enriched.get('word_count', 0):,}")
            
            with col2:
                st.metric("Sentence Count", f"{enriched.get('sentence_count', 0)}")
            
            with col3:
                st.metric("Readability", f"{enriched.get('readability_score', 0):.1f}/100")
            
            with col4:
                st.metric("Sentiment", enriched.get('sentiment', 'N/A'))
            st.write("---")
            st.write("### 🎯 Content Quality Assessment")
            
            from advanced_features import ContentQualityFilter
            quality_score = ContentQualityFilter.calculate_quality_score(doc_text, test_metadata)
            
            st.metric("Quality Score", f"{quality_score:.3f}", delta=f"{quality_score*100:.1f}%")

            if enriched.get('key_phrases'):
                st.write("### 🔑 Key Phrases")
                for i, phrase in enumerate(enriched['key_phrases'], 1):
                    st.write(f"  {i}. {phrase}")
            

            with st.expander("📋 Full Analysis"):
                st.json(enriched)
    
    else:  
        st.subheader("📋 Test with Sample Documents")
        
        sample_docs = {
            "Academic Paper": {
                "title": "Machine Learning Fundamentals",
                "content": """
                # Machine Learning Fundamentals
                
                ## Introduction
                Machine learning is a subset of artificial intelligence that enables systems to learn 
                and improve from experience without being explicitly programmed. 
                
                ## Methodology
                This study analyzes various machine learning algorithms including supervised and unsupervised learning.
                
                ## Conclusion
                The results demonstrate that ensemble methods outperform single models.
                
                ## References
                [1] Smith et al., 2020
                [2] Johnson et al., 2021
                """,
                "source": "https://example.edu/research/ml-fundamentals",
                "author": "Dr. Jane Smith",
                "publication_date": "2024-10-20"
            },
            "Blog Post": {
                "title": "Getting Started with AI",
                "content": """
                Getting Started with AI: A Beginner's Guide
                
                Hey everyone! Today I want to share my journey into AI and ML.
                
                What is AI? AI stands for Artificial Intelligence.
                
                Why learn AI? Because it's the future and it's awesome!
                
                Happy learning!
                """,
                "source": "https://example.com/blog/ai-guide",
                "author": "John Doe",
                "publication_date": "2024-10-15"
            },
            "Technical Docs": {
                "title": "API Documentation",
                "content": """
                ## REST API Documentation
                
                ### Endpoints
                - GET /api/users
                - POST /api/users
                - GET /api/users/{id}
                
                ### Example Code
                ```python
                import requests
                response = requests.get('/api/users')
                print(response.json())
                ```
                
                ### Authentication
                Use Bearer token in header.
                """,
                "source": "https://docs.example.com/api",
                "author": "Tech Team",
                "publication_date": "2024-09-01"
            }
        }
        
        selected_sample = st.selectbox("Choose a sample:", list(sample_docs.keys()))
        
        if st.button("🔍 Analyze Sample"):
            sample = sample_docs[selected_sample]
            
            st.write(f"### 📄 {sample['title']}")
            st.write(f"**Source:** {sample['source']}")
            st.write(f"**Author:** {sample['author']}")
            st.write(f"**Date:** {sample['publication_date']}")
            
            st.write("---")
            
            # Analyze with enricher
            enricher = MetadataEnricher()
            enriched = enricher.enrich_metadata(sample['content'], sample)
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Words", f"{enriched.get('word_count', 0)}")
            
            with col2:
                st.metric("Sentences", f"{enriched.get('sentence_count', 0)}")
            
            with col3:
                st.metric("Readability", f"{enriched.get('readability_score', 0):.0f}")
            
            with col4:
                st.metric("Sentiment", enriched.get('sentiment', 'N/A'))

with app_tab5:
    st.header("🔄 News/Conversation Continuity Track")
    continuity = st.session_state.continuity_system
    level1_tab, level2_tab, level3_tab, status_tab = st.tabs(
        ["📝 Level 1: Memory", "🏷️ Level 2: Topics", "🌳 Level 3: Updates", "📊 Full Status"]
    )
    
    with level1_tab:
        st.subheader("📝 Conversation Memory (Last 3 Messages)")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            user_input = st.text_area(
                "Enter message:",
                placeholder="Type a message or query...",
                height=100,
                key="level1_input"
            )
        
        with col2:
            st.write("**Quick Actions:**")
            if st.button("➕ Add Message", use_container_width=True, key="level1_add_btn"):
                if user_input:
                    continuity.process_message(user_input, role="user")
                    st.success("✅ Message added to memory!")
            
            if st.button("🧹 Clear Memory", use_container_width=True, key="level1_clear_btn"):
                continuity.level1_memory.clear()
                st.info("✅ Memory cleared!")
        
        st.write("---")
        if continuity.level1_memory.messages:
            st.write("### 💬 Recent Conversation")
            
            for i, msg in enumerate(continuity.level1_memory.messages, 1):
                color = "🟦" if msg.role == "user" else "🟩"
                
                with st.expander(f"{color} Message {i} - {msg.role.upper()}", expanded=(i==len(continuity.level1_memory.messages))):
                    st.write(f"**Content:** {msg.content}")
                    st.write(f"**Time:** {msg.timestamp}")
                    if msg.topic_tags:
                        st.write(f"**Tags:** {', '.join(msg.topic_tags)}")
            
            st.write("---")
            st.write("### 📚 Context Summary")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**Recent Context:**")
                st.code(continuity.level1_memory.get_recent_context())
            
            with col2:
                st.write("**Previous Summary:**")
                summary = continuity.level1_memory.get_context_summary()
                if summary != "No previous context":
                    st.code(summary)
                else:
                    st.info("No previous messages yet")
        else:
            st.info("📭 No messages yet. Add a message to start!")
    
    with level2_tab:
        st.subheader("🏷️ Topic Tracking & Auto-Detection")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            topic_input = st.text_area(
                "Enter text to track topic:",
                placeholder="Type any text to auto-detect or create topic...",
                height=100,
                key="level2_input"
            )
        
        with col2:
            st.write("**Topic Actions:**")
            if st.button("🔎 Detect Topic", use_container_width=True, key="level2_detect_btn"):
                if topic_input:
                    topic_name, similarity = continuity.level2_tracker.detect_or_create_topic(topic_input)
                    st.success(f"✅ Topic: **{topic_name}**")
                    st.metric("Similarity Score", f"{similarity:.2%}")
        
        st.write("---")
        topics = continuity.level2_tracker.get_all_topics()
        
        if topics:
            st.write(f"### 📊 All Topics ({len(topics)})")
            
            for topic in topics:
                with st.expander(f"**{topic['name']}** ({topic['message_count']} messages)", expanded=False):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("**Created:** " + str(topic['created']))
                        st.write("**Message Count:** " + str(topic['message_count']))
                    
                    with col2:
                        st.write("**Recent Messages:**")
                        if topic.get('recent_messages'):
                            for msg in topic['recent_messages'][:3]:
                                st.write(f"• {msg[:100]}...")
                    
                    detailed_topic = continuity.level2_tracker.get_topic_by_name(topic['name'])
                    if detailed_topic:
                        st.write("**Keywords:**")
                        if detailed_topic.get('keywords'):
                            st.write(", ".join(detailed_topic['keywords'][:10]))
        else:
            st.info("🏷️ No topics yet. Add messages to auto-create topics!")
    
    with level3_tab:
        st.subheader("🌳 Conversation Branching & Merging")
        
        st.write("**Append updates to existing topics or create alternative branches:**")

        topics = continuity.level2_tracker.get_all_topics()
        topic_names = [t['name'] for t in topics]
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("### ➕ Append Update")
            
            if topic_names:
                selected_topic = st.selectbox("Select topic:", topic_names, key="branch_append_topic")
                append_message = st.text_area("New information:", height=100, key="append_msg")
                
                if st.button("📌 Append to Topic", use_container_width=True, key="append_btn"):
                    if append_message:
                        result = continuity.level3_updater.append_to_conversation(
                            selected_topic, append_message, "update"
                        )
                        st.success(f"✅ {result['status']}")
            else:
                st.info("Create topics first in Level 2")
        
        with col2:
            st.write("### 🌳 Create Branch")
            
            if topic_names:
                branch_topic = st.selectbox("Branch from topic:", topic_names, key="branch_topic")
                branch_name = st.text_input("Branch name:", placeholder="e.g., Best Case", key="branch_name")
                branch_content = st.text_area("Branch content:", height=100, key="branch_content")
                
                if st.button("🔀 Create Branch", use_container_width=True, key="create_branch_btn"):
                    if branch_name and branch_content:
                        result = continuity.level3_updater.branch_conversation(
                            branch_topic, branch_name, branch_content
                        )
                        st.success(f"✅ Branch created: {branch_name}")
                        st.balloons()
                    else:
                        st.warning("⚠️ Please enter both branch name and content")
            else:
                st.info("Create topics first in Level 2")
        
        st.write("---")
        st.write("### 🔄 Conversation Evolution")
        
        if topic_names:
            selected_topic_view = st.selectbox("View evolution of:", topic_names, key="view_evolution")
            evolution = continuity.level3_updater.get_conversation_evolution(selected_topic_view)
            
            if evolution.get('branches'):
                st.write(f"**Branches ({len(evolution['branches'])}):**")
                for branch in evolution['branches']:
                    with st.expander(f"🔀 {branch['branch_name']}"):
                        st.write(branch['content'])
            else:
                st.info("No branches yet. Create one above!")
        else:
            st.info("No topics to view")
    
    with status_tab:
        st.subheader("📊 Complete System Status")
        
        status = continuity.get_full_status()
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Recent Messages", len(status['level1']['recent_messages']))
        
        with col2:
            st.metric("Total Topics", status['level2']['total_topics'])
        
        with col3:
            st.metric("Active Branches", status['level3']['total_branches'])
        
        with col4:
            st.metric("Updates Logged", len(status['level3']['update_timeline']))
        
        st.write("---")

        st.write("### 📝 Level 1: Memory Status")
        
        if status['level1']['recent_messages']:
            for i, msg in enumerate(status['level1']['recent_messages'], 1):
                st.write(f"**{i}. {msg['role'].upper()}:** {msg['content'][:100]}...")
            
            st.write("**Context Summary:**")
            st.code(status['level1']['context_summary'])
        else:
            st.info("No messages in memory")
        
        st.write("---")
        st.write("### 🏷️ Level 2: Topic Status")
        
        topics_list = status['level2']['topics']
        if topics_list:
            topic_data = {
                'Topic': [t['name'] for t in topics_list],
                'Messages': [t['message_count'] for t in topics_list],
                'Created': [t['created'] for t in topics_list]
            }
            st.table(topic_data)
        else:
            st.info("No topics created yet")
        
        st.write("---")
        st.write("### 🌳 Level 3: Update Timeline")
        
        updates = status['level3']['update_timeline']
        if updates:
            for update in updates[:10]:
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.write(f"**{update['action'].upper()}** - {update['topic']}")
                    if update['action'] == 'branch':
                        st.write(f"└─ Branch: {update['branch_name']}")
                    st.caption(update['timestamp'])
                
                with col2:
                    status_badge = f"✅ {update['status']}"
                    st.write(status_badge)
        else:
            st.info("No updates yet")
        
        st.write("---")
        
        # Export status
        st.write("### 📥 Export Status")
        
        if st.button("📊 Export as JSON", key="export_json_btn"):
            import json
            json_str = json.dumps(status, indent=2, default=str)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"conversation_status_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
